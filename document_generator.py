import os
import json
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Класс для генерации документов в формате Word (DOCX)"""

    def __init__(self, app=None):
        self.app = app
        self.output_dir = None
        self.protocols_dir = None
        self.resolutions_dir = None

        if app:
            self.init_app(app)

    def init_app(self, app):
        """Инициализация с приложением Flask"""
        self.output_dir = app.config['OUTPUT_DIR']
        self.protocols_dir = app.config['PROTOCOLS_DIR']
        self.resolutions_dir = app.config['RESOLUTIONS_DIR']

        # Создаем папки, если их нет
        os.makedirs(self.protocols_dir, exist_ok=True)
        os.makedirs(self.resolutions_dir, exist_ok=True)
        logger.info(f"Директории для документов созданы: {self.output_dir}")

    def _setup_document_styles(self, doc):
        """Настройка стилей для документа"""
        try:
            # Настраиваем нормальный стиль
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(14)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.first_line_indent = Cm(1.25)

            # Создаем стиль для заголовка если его нет
            if 'Title' not in [s.name for s in doc.styles]:
                title_style = doc.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
                title_style.font.name = 'Times New Roman'
                title_style.font.size = Pt(18)
                title_style.font.bold = True
                title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_style.paragraph_format.space_after = Pt(12)
                title_style.paragraph_format.space_before = Pt(12)

            # Создаем стиль для подзаголовка
            if 'Subtitle' not in [s.name for s in doc.styles]:
                subtitle_style = doc.styles.add_style('Subtitle', WD_STYLE_TYPE.PARAGRAPH)
                subtitle_style.font.name = 'Times New Roman'
                subtitle_style.font.size = Pt(16)
                subtitle_style.font.bold = True
                subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle_style.paragraph_format.space_after = Pt(10)

            # Настраиваем стиль заголовка 1
            heading1 = doc.styles['Heading 1']
            heading1.font.name = 'Times New Roman'
            heading1.font.size = Pt(16)
            heading1.font.bold = True
            heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading1.paragraph_format.space_before = Pt(12)
            heading1.paragraph_format.space_after = Pt(6)

        except Exception as e:
            logger.warning(f"Ошибка при настройке стилей: {e}")

    def generate_protocol(self, data):
        """Генерация протокола в формате Word"""
        try:
            # Формируем имя файла
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"Протокол_{data.get('number', '001')}_{timestamp}.docx"
            filepath = os.path.join(self.protocols_dir, filename)

            # Создаем Word документ
            doc = DocxDocument()

            # Настраиваем стили
            self._setup_document_styles(doc)

            # Заголовок
            title = doc.add_heading('ПРОТОКОЛ', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Номер и дата
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'№ {data.get("number", "")} от {data.get("date", "")}')
            run.bold = True
            run.font.size = Pt(14)

            # Место составления
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('г. Москва')

            doc.add_paragraph()  # Пустая строка

            # Тема
            p = doc.add_paragraph()
            run = p.add_run('Тема: ')
            run.bold = True
            p.add_run(data.get('topic', ''))

            # Место проведения
            p = doc.add_paragraph()
            run = p.add_run('Место проведения: ')
            run.bold = True
            p.add_run(data.get('location', ''))

            # Дата и время
            p = doc.add_paragraph()
            run = p.add_run('Дата и время: ')
            run.bold = True
            p.add_run(data.get('datetime', ''))

            doc.add_paragraph()  # Пустая строка

            # Председатель и секретарь
            if data.get('chairman'):
                p = doc.add_paragraph()
                run = p.add_run('Председатель: ')
                run.bold = True
                p.add_run(data.get('chairman', ''))

            if data.get('secretary'):
                p = doc.add_paragraph()
                run = p.add_run('Секретарь: ')
                run.bold = True
                p.add_run(data.get('secretary', ''))

            doc.add_paragraph()

            # Присутствовали
            participants = data.get('participants', '').split('\n')
            participants = [p.strip() for p in participants if p.strip()]
            if participants:
                p = doc.add_paragraph()
                run = p.add_run('Присутствовали:')
                run.bold = True
                for participant in participants:
                    p = doc.add_paragraph(f'• {participant}', style='List Bullet')
                doc.add_paragraph()

            # Повестка дня
            agenda = data.get('agenda', '').split('\n')
            agenda = [a.strip() for a in agenda if a.strip()]
            if agenda:
                p = doc.add_paragraph()
                run = p.add_run('Повестка дня:')
                run.bold = True
                for i, item in enumerate(agenda, 1):
                    p = doc.add_paragraph(f'{i}. {item}', style='List Number')
                doc.add_paragraph()

            # Решения
            decisions = data.get('decisions', '').split('\n')
            decisions = [d.strip() for d in decisions if d.strip()]
            if decisions:
                p = doc.add_paragraph()
                run = p.add_run('Постановили:')
                run.bold = True
                for i, decision in enumerate(decisions, 1):
                    p = doc.add_paragraph(f'{i}. {decision}', style='List Number')
                doc.add_paragraph()

            # Подписи
            doc.add_paragraph()  # Пустая строка
            doc.add_paragraph()  # Пустая строка

            p = doc.add_paragraph()
            p.add_run('Председатель:').bold = True
            p.add_run(f' __________________ {data.get("chairman", "")}')

            p = doc.add_paragraph()
            p.add_run('Секретарь:').bold = True
            p.add_run(f' __________________ {data.get("secretary", "")}')

            # Сохраняем документ
            doc.save(filepath)
            logger.info(f"Word документ создан: {filepath}")

            return filepath

        except Exception as e:
            logger.error(f"Ошибка при создании Word документа: {e}")
            raise

    def generate_resolution(self, data):
        """Генерация постановления в формате Word"""
        try:
            # Формируем имя файла
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"Постановление_{data.get('number', '001')}_{timestamp}.docx"
            filepath = os.path.join(self.resolutions_dir, filename)

            # Создаем Word документ
            doc = DocxDocument()

            # Настраиваем стили
            self._setup_document_styles(doc)

            # Заголовок
            title = doc.add_heading('ПОСТАНОВЛЕНИЕ', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Номер и дата
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'№ {data.get("number", "")} от {data.get("date", "")}')
            run.bold = True
            run.font.size = Pt(14)

            # Место составления
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('г. Москва')

            doc.add_paragraph()  # Пустая строка

            # Тема
            p = doc.add_paragraph()
            run = p.add_run('Тема: ')
            run.bold = True
            p.add_run(data.get('topic', ''))

            doc.add_paragraph()

            # Основание
            if data.get('basis'):
                p = doc.add_paragraph()
                run = p.add_run('Основание:')
                run.bold = True
                p = doc.add_paragraph(data.get('basis', ''))
                doc.add_paragraph()

            # Постановляющая часть
            p = doc.add_paragraph()
            run = p.add_run('ПОСТАНОВЛЯЕТ:')
            run.bold = True
            run.font.size = Pt(14)

            doc.add_paragraph()

            # Текст постановления
            resolutions = data.get('text', '').split('\n')
            resolutions = [r.strip() for r in resolutions if r.strip()]
            if resolutions:
                for i, resolution in enumerate(resolutions, 1):
                    p = doc.add_paragraph(f'{i}. {resolution}', style='List Number')
                doc.add_paragraph()

            # Контроль исполнения
            p = doc.add_paragraph(
                '3. Контроль за исполнением настоящего постановления оставляю за собой.',
                style='List Number'
            )

            doc.add_paragraph()
            doc.add_paragraph()

            # Подписи
            if data.get('chairman'):
                p = doc.add_paragraph()
                p.add_run('Председатель:').bold = True
                p.add_run(f' __________________ {data.get("chairman", "")}')
                doc.add_paragraph()

            # Члены комиссии
            members = data.get('members', '').split('\n')
            members = [m.strip() for m in members if m.strip()]
            for i, member in enumerate(members, 1):
                p = doc.add_paragraph()
                p.add_run(f'Член комиссии {i}:').bold = True
                p.add_run(f' __________________ {member}')

            # Печать
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('М.П.').bold = True

            # Сохраняем документ
            doc.save(filepath)
            logger.info(f"Word документ создан: {filepath}")

            return filepath

        except Exception as e:
            logger.error(f"Ошибка при создании Word документа: {e}")
            raise


# Создаем экземпляр генератора
document_generator = DocumentGenerator()